"""Actualiza el .docx de la entrega de la Actividad 3.1 con:
    - Matriz de confusión real del SVM.
    - Subsección sobre el Supervisor de barriles.
    - Descripción actualizada de la detección con Sliding Window Search.
    - Bloque de código reemplazado por el vehicle_controller.py actual.
    - Marcador para el enlace de YouTube.
    - Declaración explícita del uso de herramientas de IA.

Idempotente: se puede correr varias veces; vuelve a generar las secciones
modificadas a partir de marcadores fijos en el documento.
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# ---------------------------------------------------------------------------
# Rutas.
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent
DOCX_PATH = (
    Path("/Users/armand501/Library/Mobile Documents/com~apple~CloudDocs/")
    / "My Documents/Maestria/Navegación autónoma "
    / "Actividad_3.1_Detección_de_Peatones_Equipo19.docx"
)
DOCX_BACKUP = DOCX_PATH.with_name(DOCX_PATH.stem + "_backup_pre_update.docx")
CONTROLLER_PATH = (
    REPO_ROOT
    / "SDC_webots/controllers/vehicle_controller/vehicle_controller.py"
)

# ---------------------------------------------------------------------------
# Datos a inyectar.
# ---------------------------------------------------------------------------
CONFUSION_MATRIX_LINES = [
    "Matriz obtenida con un split 80/20 sobre 659 muestras (420 positivas, 239 negativas):",
    "",
    "                     Predicho",
    "                   No peatón   Peatón",
    "    Real No peatón     47        1",
    "    Real Peatón         3       81",
    "",
    "Reporte de clasificación:",
    "                  precision   recall   f1-score   support",
    "    No peatón        0.94      0.98       0.96        48",
    "    Peatón           0.99      0.96       0.98        84",
    "    accuracy                              0.97       132",
    "",
    "Exactitud total del modelo: 0.97 sobre el conjunto de prueba.",
]

SUPERVISOR_SUBSECTION_TITLE = "5. Supervisor de aparición de barriles"
SUPERVISOR_SUBSECTION_BODY = [
    "El mundo de la actividad incluye un robot Supervisor cuyo controlador "
    "(supervisor_controller.py, provisto por el profesor y utilizado sin "
    "modificaciones) coloca periódicamente un barril DEF BARREL frente al "
    "vehículo y lo elimina segundos después. El controlador del vehículo "
    "no tiene que invocar a este supervisor: simplemente reacciona ante el "
    "obstáculo cuando aparece dentro del rango del LiDAR.",
]

UPDATED_SVM_DESCRIPTION = [
    "Se exportó el modelo previamente entrenado:",
    "svm_pedestrian_model.pkl",
    "El modelo fue cargado dentro del controlador utilizando la librería joblib.",
    "model = joblib.load(\"svm_pedestrian_model.pkl\")",
    "Para realizar la detección de peatones se utilizó:",
    "• OpenCV",
    "• HOG Descriptor (ventana de 64x128)",
    "• SVM lineal entrenado con Scikit-Learn",
    "",
    "Las muestras de entrenamiento se obtuvieron del dataset Penn-Fudan "
    "Pedestrian Database. Las MUESTRAS POSITIVAS se generaron recortando "
    "cada bounding box anotado en los archivos PennFudanPed/Annotation/*.txt, "
    "de manera que el SVM aprendiera la apariencia del peatón y no del fondo. "
    "Las MUESTRAS NEGATIVAS se generaron recortando regiones aleatorias de las "
    "mismas imágenes que no se solaparan con ningún bounding box positivo.",
    "",
    "Debido a que la cámara del mundo de Webots trabaja con resolución 256x128 "
    "y no coincide con la ventana de entrenamiento del SVM (64x128), se aplicó "
    "la técnica de Sliding Window Search multi-escala sobre la región de "
    "interés inferior de la imagen, tal como recomienda la actividad. Para "
    "cada frame se evalúan ~63 ventanas a 3 escalas distintas y la predicción "
    "se hace en batch con LinearSVC.predict para mantener bajo el costo por "
    "paso de simulación (32 ms).",
]

YOUTUBE_PLACEHOLDER = (
    "[PEGAR ENLACE DE YOUTUBE AQUÍ ANTES DE LA ENTREGA]"
)

AI_DECLARATION_TITLE = "Declaración de uso de inteligencia artificial"
AI_DECLARATION_BODY = [
    "De acuerdo con los lineamientos de la actividad, se declara el uso "
    "de herramientas de inteligencia artificial durante el desarrollo:",
    "",
    "Anthropic. (2026). Claude Opus 4.7 [Modelo de lenguaje grande]. "
    "Utilizado para generación de código (controlador, script de "
    "entrenamiento del SVM, launcher), depuración (corrección de la "
    "orquestación entre Webots y el controlador externo, manejo de "
    "puertos y cierre limpio de la simulación) y optimización (Sliding "
    "Window Search con predicción en batch). https://www.anthropic.com/claude",
    "",
    "La responsabilidad final sobre el contenido entregado recae en los "
    "autores. Las soluciones presentadas fueron revisadas, ajustadas y "
    "validadas por el equipo en Webots.",
]


# ---------------------------------------------------------------------------
# Helpers para manipular el XML del documento.
# ---------------------------------------------------------------------------
def find_paragraph_index(doc, text_starts_with: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(text_starts_with):
            return i
    raise ValueError(f"No se encontró un párrafo que empiece con: {text_starts_with!r}")


def delete_paragraphs_between(doc, start_exclusive: int, end_exclusive: int) -> None:
    """Borra los párrafos cuyos índices están en (start, end) exclusivos."""
    paragraphs = doc.paragraphs[start_exclusive + 1 : end_exclusive]
    for p in paragraphs:
        p._element.getparent().remove(p._element)


def delete_paragraph_range(doc, start_inclusive: int, end_exclusive: int) -> None:
    """Borra los párrafos cuyos índices están en [start, end)."""
    paragraphs = doc.paragraphs[start_inclusive:end_exclusive]
    for p in paragraphs:
        p._element.getparent().remove(p._element)


def insert_paragraph_after(target_paragraph, text: str, style_paragraph=None):
    """Inserta un nuevo párrafo justo después de target_paragraph y devuelve el objeto Paragraph."""
    new_p = deepcopy((style_paragraph or target_paragraph)._element)
    # Vaciar el contenido clonado.
    for child in list(new_p):
        new_p.remove(child)
    # Insertar después del target en el árbol XML.
    target_paragraph._element.addnext(new_p)
    # Envolver con la clase Paragraph de python-docx.
    from docx.text.paragraph import Paragraph

    paragraph = Paragraph(new_p, target_paragraph._parent)
    if text:
        run = paragraph.add_run(text)
    return paragraph


def insert_block_after(target_paragraph, lines, style_paragraph=None, monospace=False):
    """Inserta una lista de líneas como párrafos consecutivos después de target."""
    last = target_paragraph
    for line in lines:
        new_p = insert_paragraph_after(last, line, style_paragraph=style_paragraph)
        if monospace:
            for run in new_p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        last = new_p
    return last


# ---------------------------------------------------------------------------
# Marcadores idempotentes — si ya se inyectaron antes, se eliminan primero.
# ---------------------------------------------------------------------------
SECTION_MARKERS = {
    "supervisor_subsection_start": SUPERVISOR_SUBSECTION_TITLE,
    "ai_declaration_title": AI_DECLARATION_TITLE,
}


def remove_section_if_present(doc, marker_text: str, terminator_text: str) -> None:
    """Si existe un párrafo con marker_text, borra desde ahí (inclusive) hasta
    el siguiente párrafo cuyo texto empiece con terminator_text (exclusivo)."""
    paragraphs = doc.paragraphs
    start = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == marker_text:
            start = i
            break
    if start is None:
        return
    end = len(paragraphs)
    for j in range(start + 1, len(paragraphs)):
        if paragraphs[j].text.strip().startswith(terminator_text):
            end = j
            break
    delete_paragraph_range(doc, start, end)


# ---------------------------------------------------------------------------
# Edición principal.
# ---------------------------------------------------------------------------
def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"No se encontró el docx: {DOCX_PATH}")
    if not CONTROLLER_PATH.exists():
        raise SystemExit(f"No se encontró el controlador: {CONTROLLER_PATH}")

    # Backup antes de tocar nada.
    shutil.copyfile(DOCX_PATH, DOCX_BACKUP)
    print(f"[backup] {DOCX_BACKUP}")

    doc = Document(str(DOCX_PATH))
    controller_code = CONTROLLER_PATH.read_text(encoding="utf-8")

    # Estilo de párrafo de referencia (uno cualquiera del cuerpo).
    body_style = doc.paragraphs[12]  # "Introducción"

    # 0) Limpieza idempotente de secciones que vamos a re-generar.
    remove_section_if_present(
        doc,
        SUPERVISOR_SUBSECTION_TITLE,
        terminator_text="Resultados obtenidos",
    )
    remove_section_if_present(
        doc,
        AI_DECLARATION_TITLE,
        terminator_text="Conclusiones",
    )

    # 1) Actualizar la sección 2 (descripción del SVM).
    sec2_idx = find_paragraph_index(doc, "2. Integración del modelo SVM")
    sec3_idx = find_paragraph_index(doc, "3. Integración del controlador PID")
    # Borrar contenido entre el título y el siguiente título (exclusivos).
    delete_paragraphs_between(doc, sec2_idx, sec3_idx)
    # Re-insertar la descripción nueva.
    sec2_paragraph = doc.paragraphs[sec2_idx]
    insert_block_after(sec2_paragraph, UPDATED_SVM_DESCRIPTION, style_paragraph=body_style)

    # 2) Insertar subsección del supervisor justo antes de "Resultados obtenidos".
    resultados_idx = find_paragraph_index(doc, "Resultados obtenidos")
    resultados_paragraph = doc.paragraphs[resultados_idx]
    # Insertamos en orden inverso para que aparezcan en el orden correcto.
    # Estrategia: insertar el bloque después del párrafo previo a "Resultados".
    previous_paragraph = doc.paragraphs[resultados_idx - 1]
    # Título primero.
    title_p = insert_paragraph_after(
        previous_paragraph, SUPERVISOR_SUBSECTION_TITLE, style_paragraph=body_style
    )
    for run in title_p.runs:
        run.bold = True
    insert_block_after(title_p, SUPERVISOR_SUBSECTION_BODY, style_paragraph=body_style)

    # 3) Inyectar la matriz de confusión después de su título.
    cm_idx = find_paragraph_index(doc, "Matriz de confusión del modelo SVM")
    cm_paragraph = doc.paragraphs[cm_idx]
    # Borrar cualquier contenido viejo entre el título y "Script final".
    script_idx = find_paragraph_index(doc, "Script final del controlador")
    delete_paragraphs_between(doc, cm_idx, script_idx)
    insert_block_after(cm_paragraph, CONFUSION_MATRIX_LINES, style_paragraph=body_style, monospace=True)

    # 4) Reemplazar el bloque de código.
    codigo_idx = find_paragraph_index(doc, "Código completo")
    video_idx = find_paragraph_index(doc, "Video de demostración")
    delete_paragraphs_between(doc, codigo_idx, video_idx)
    codigo_paragraph = doc.paragraphs[codigo_idx]
    insert_block_after(
        codigo_paragraph,
        controller_code.splitlines(),
        style_paragraph=body_style,
        monospace=True,
    )

    # 5) Reemplazar el placeholder de YouTube.
    youtube_idx = find_paragraph_index(doc, "Enlace de YouTube")
    conclusiones_idx = find_paragraph_index(doc, "Conclusiones")
    delete_paragraphs_between(doc, youtube_idx, conclusiones_idx)
    youtube_paragraph = doc.paragraphs[youtube_idx]
    insert_paragraph_after(youtube_paragraph, YOUTUBE_PLACEHOLDER, style_paragraph=body_style)

    # 6) Insertar declaración de IA justo antes de "Bibliografía".
    bibliografia_idx = find_paragraph_index(doc, "Bibliografía")
    biblio_prev = doc.paragraphs[bibliografia_idx - 1]
    ai_title = insert_paragraph_after(biblio_prev, AI_DECLARATION_TITLE, style_paragraph=body_style)
    for run in ai_title.runs:
        run.bold = True
    insert_block_after(ai_title, AI_DECLARATION_BODY, style_paragraph=body_style)

    # Guardar.
    doc.save(str(DOCX_PATH))
    print(f"[done] Actualizado: {DOCX_PATH}")


if __name__ == "__main__":
    main()
